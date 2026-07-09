"""
Filtro automático de roles sugeridos a partir del texto del CV, por
coincidencia de palabras clave (sin llamadas a ninguna API de IA).

Este módulo es una ayuda para el admin (primer filtro rápido apenas se
sube el CV) — no reemplaza el análisis de `cv_scores.roles_objetivo` que
arma el admin a mano y que sí ve el candidato final.
"""
import io
import re
import unicodedata

import docx
import pdfplumber

from services.role_keywords import ROLE_KEYWORDS


def extract_text(content: bytes, ext: str) -> str:
    """Extrae el texto plano de un CV en PDF, DOCX o TXT. Para .doc (binario
    legado) o formatos no soportados devuelve "" — no hay parser puro-Python
    confiable sin agregar dependencias de sistema (ej. pandoc)."""
    if ext == ".pdf":
        parts = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                parts.append(page.extract_text() or "")
        return "\n".join(parts)

    if ext == ".docx":
        document = docx.Document(io.BytesIO(content))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)

    if ext == ".txt":
        return content.decode("utf-8", errors="ignore")

    return ""


def normalize_text(text: str) -> str:
    """Minúsculas + sin tildes/diacríticos + espacios colapsados."""
    text = text.lower()
    text = unicodedata.normalize("NFKD", text)
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    return re.sub(r"\s+", " ", text).strip()


def _keyword_pattern(keyword: str) -> re.Pattern:
    return re.compile(r"\b" + re.escape(normalize_text(keyword)) + r"\b")


def match_roles(text: str, top_n: int = 5, min_score: float = 0.15) -> list[dict]:
    """Puntúa cada puesto de ROLE_KEYWORDS contra el texto normalizado del
    CV. Score = fracción de keywords de ese puesto encontradas (presencia,
    no frecuencia). Devuelve el top_n con score >= min_score, ordenado
    descendente (empate alfabético)."""
    normalized = normalize_text(text)
    if not normalized:
        return []

    results = []
    for titulo, data in ROLE_KEYWORDS.items():
        keywords = data["keywords"]
        weight = data.get("weight", 1.0)
        if not keywords:
            continue

        encontrados = [kw for kw in keywords if _keyword_pattern(kw).search(normalized)]
        raw_score = (len(encontrados) / len(keywords)) * weight
        if raw_score < min_score:
            continue

        results.append({
            "titulo": titulo,
            "match_porcentaje": min(round(raw_score * 100), 100),
            "keywords_encontrados": encontrados,
        })

    results.sort(key=lambda r: (-r["match_porcentaje"], r["titulo"]))
    return results[:top_n]


def suggest_roles_from_cv(content: bytes, ext: str, top_n: int = 5) -> list[dict]:
    """Extrae el texto del CV y devuelve los puestos sugeridos. Puede lanzar
    excepción si la extracción falla (PDF corrupto, etc.) — el caller debe
    contenerla para no romper el flujo principal de subida del CV."""
    text = extract_text(content, ext)
    return match_roles(text, top_n=top_n)
